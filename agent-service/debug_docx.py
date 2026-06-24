"""Debug DOCX encoding"""
from docx import Document

doc = Document('/tmp/test_chinese.docx')
text = doc.paragraphs[0].text
print('Raw text:', repr(text))
print('Hex:', text.encode('utf-8', errors='replace').hex()[:80])
print('Expected "测":', '测'.encode('utf-8').hex())

p1 = doc.paragraphs[1].text
print('Para 1:', repr(p1))

# Check: write to file to verify encoding
with open('/tmp/debug_output.txt', 'w', encoding='utf-8') as f:
    for p in doc.paragraphs:
        f.write(p.text + '\n')
print('Written to /tmp/debug_output.txt')
print('File size:', end=' ')
import os
print(os.path.getsize('/tmp/debug_output.txt'))