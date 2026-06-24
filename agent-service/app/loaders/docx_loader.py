"""DOCX 加载器 — 提取段落、标题、表格、页眉、页脚，支持 XML 回退"""

import logging
import zipfile

logger = logging.getLogger(__name__)

# OOXML 命名空间
NSMAP = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006",
}


class DocxLoader:
    """读取 .docx 全部文本内容，多层回退确保最大程度提取"""

    @staticmethod
    def supports(extension: str) -> bool:
        return extension.lower() in ("docx",)

    def load(self, file_path: str) -> str:
        from docx import Document

        logger.info("[DocxLoader] 开始解析: %s", file_path)

        # ── 第一层：python-docx 标准解析 ──
        doc = Document(file_path)
        parts = []

        headers_text = self._extract_headers(doc)
        if headers_text.strip():
            parts.append(headers_text)
            parts.append("")

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                parts.append("")
                continue

            heading_level = self._get_heading_level(para)
            if heading_level > 0:
                prefix = "#" * min(heading_level, 3)
                parts.append(f"{prefix} {text}")
            else:
                parts.append(text)

        for table in doc.tables:
            parts.append("")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                parts.append(" | ".join(cells))
            parts.append("")

        footers_text = self._extract_footers(doc)
        if footers_text.strip():
            parts.append("")
            parts.append(footers_text)

        text = "\n".join(parts)
        logger.info("[DocxLoader] 标准解析: 段落数=%d, 表格数=%d, 文本长度=%d",
                     len(doc.paragraphs), len(doc.tables), len(text.strip()))

        # ── 第二层：如果标准解析结果为空，尝试原始 XML 回退 ──
        if not text.strip():
            logger.warning("[DocxLoader] 标准解析无文本内容，尝试 XML 回退提取")
            xml_text = self._extract_from_raw_xml(file_path)
            if xml_text.strip():
                parts = [xml_text]
                text = xml_text
                logger.info("[DocxLoader] XML 回退提取成功，文本长度=%d", len(xml_text.strip()))

        # ── 第三层：如果仍然为空，检测是否为纯图片文档 ──
        if not text.strip():
            drawing_count = self._count_drawings(file_path)
            if drawing_count > 0:
                text = f"[该文档为纯图片文档，包含 {drawing_count} 张图片，无可提取的文字内容]"
                logger.info("[DocxLoader] 检测到纯图片文档，图片数=%d", drawing_count)
            else:
                text = "[该文档无可提取的文字内容]"
                logger.info("[DocxLoader] 文档无任何可提取内容")

        logger.info("[DocxLoader] 最终解析完成, 总长度=%d", len(text))
        return text

    def _extract_from_raw_xml(self, file_path: str) -> str:
        """原始 XML 回退：绕过 python-docx API，直接从 document.xml 提取 <w:t> 文本。
        用于处理 python-docx 无法访问的内容（如结构化文档标签 SDT、文本框等）。"""
        try:
            from lxml import etree

            with zipfile.ZipFile(file_path) as zf:
                if "word/document.xml" not in zf.namelist():
                    return ""
                xml_bytes = zf.read("word/document.xml")

            root = etree.fromstring(xml_bytes)

            # 处理 mc:AlternateContent — 选择 Fallback 分支以保证兼容性
            for ac in root.findall(".//mc:AlternateContent", NSMAP):
                fallback = ac.find("./mc:Fallback", NSMAP)
                if fallback is not None:
                    parent = ac.getparent()
                    if parent is not None:
                        idx = list(parent).index(ac)
                        parent.remove(ac)
                        for child in fallback:
                            parent.insert(idx, child)
                            idx += 1

            texts = []
            for t_elem in root.findall(".//w:t", NSMAP):
                if t_elem.text and t_elem.text.strip():
                    # 检查父元素是否有 xml:space="preserve"
                    preserve_space = t_elem.get("{http://www.w3.org/XML/1998/namespace}space") == "preserve"
                    if preserve_space:
                        texts.append(t_elem.text)
                    else:
                        texts.append(t_elem.text.strip())

            return "\n".join(texts)

        except Exception as e:
            logger.warning("[DocxLoader] XML 回退提取失败: %s", e)
            return ""

    def _count_drawings(self, file_path: str) -> int:
        """统计文档中的图片/绘图数量"""
        try:
            from lxml import etree

            with zipfile.ZipFile(file_path) as zf:
                if "word/document.xml" not in zf.namelist():
                    return 0
                xml_bytes = zf.read("word/document.xml")

            root = etree.fromstring(xml_bytes)
            drawings = root.findall(".//w:drawing", NSMAP)
            return len(drawings)

        except Exception:
            return 0

    def _get_heading_level(self, para) -> int:
        try:
            style = para.style
            if style is None:
                return 0
            name = style.name
            if name is None:
                return 0
            if name.startswith("Heading") or name.startswith("heading"):
                level_str = name.replace("Heading", "").replace("heading", "").strip()
                if level_str.isdigit():
                    return int(level_str)
                return 1
            if name.startswith("Title") or name.startswith("Subtitle"):
                return 1
        except Exception:
            pass
        return 0

    def _extract_headers(self, doc) -> str:
        parts = []
        for section in doc.sections:
            try:
                header = section.header
                if header and not header.is_linked_to_previous:
                    for para in header.paragraphs:
                        text = para.text.strip()
                        if text:
                            parts.append("[页眉] " + text)
            except Exception:
                pass
        return "\n".join(parts) if parts else ""

    def _extract_footers(self, doc) -> str:
        parts = []
        for section in doc.sections:
            try:
                footer = section.footer
                if footer and not footer.is_linked_to_previous:
                    for para in footer.paragraphs:
                        text = para.text.strip()
                        if text:
                            parts.append("[页脚] " + text)
            except Exception:
                pass
        return "\n".join(parts) if parts else ""